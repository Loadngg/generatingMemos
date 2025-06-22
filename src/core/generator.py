import os
from typing import Callable, List

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

from src.core import utils
from src.core.models.field import Field
from src.core.models.record import Record


class Generator:
    doc: DocumentObject = None

    def generate(self, template_path: str, records: List[Record], output_path: str,
                 update_info_label: Callable, enable_generating: Callable) -> None:
        last_generated_name: str = ''

        output_path = output_path + template_path.split('/')[-1].split('.')[0] + '/'
        if not os.path.exists(output_path):
            os.mkdir(output_path)

        for item in records:
            if not item.full_name or not item.full_name.text:
                raise
            second_name = item.full_name.text.split(' ')[0]
            if last_generated_name != item.full_name.text:
                last_generated_name = item.full_name.text

            self.doc = Document(template_path)
            self._set_content(item)

            title = ""
            if item.event_title and item.event_title.text:
                title = item.event_title.text
            else:
                titles = []
                for data in item.events.values():
                    titles.append(data[1].text)

                title = " ".join(titles)

            title = " ".join(title.split()).strip()

            self._save(second_name, output_path, title)
            update_info_label(f"сгенерировано для {second_name}")

        update_info_label(f'файлы были успешно сгенерированы - {len(records)}')
        enable_generating()

    def _set_institute(self, institute: Field) -> None:
        table_text: str = self.doc.tables[0].rows[0].cells[1].text
        self.doc.tables[0].rows[0].cells[1].text = table_text.replace(institute.tag, institute.text)
        self.doc.tables[0].rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _set_content(self, item: Record) -> None:
        item_fields = item.get_fields()

        self._set_institute(item.institute)

        for p in self.doc.paragraphs:
            for field in item_fields:
                if type(field) == Field:
                    self._set_text(p, field)
                elif type(field) == dict:
                    for data in field.values():
                        for data_field in data:
                            self._set_text(p, data_field)

    @staticmethod
    def _set_text(p: Paragraph, field: Field) -> None:
        key = field.tag
        value = field.text
        if value is None:
            return
        p.text = p.text.replace(key, value)

    def _save(self, second_name: str, output_path: str, title: str) -> None:
        output_title = ' '.join(title.split(' ')[:5])
        output_name = (f"{utils.get_translit(second_name)} "
                       f"{utils.get_translit(output_title)}")

        output_file = output_path + output_name + '.docx'
        pdf_path = output_path + second_name

        self.doc.save(output_file)
        utils.save_pdf(output_file, pdf_path, output_name)
